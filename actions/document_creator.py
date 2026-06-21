"""document_creator.py — Create Word, Excel, Google Docs, Google Sheets."""
from __future__ import annotations
import os
from pathlib import Path


def document_creator(parameters=None, player=None, **kwargs):
    if parameters is None:
        parameters = {}
    action = parameters.get("action", "word")
    title = parameters.get("title", "Documento")
    content = parameters.get("content", "")
    sheets = parameters.get("sheets", [])
    save_path = parameters.get("save_path", str(Path.home() / "Documents"))

    os.makedirs(save_path, exist_ok=True)

    if action == "word":
        try:
            from docx import Document
        except ImportError:
            return "python-docx no instalado. Ejecutá: pip install python-docx"
        doc = Document()
        doc.add_heading(title, 0)
        for line in content.split("\n"):
            line = line.strip()
            if line.startswith("## "):
                doc.add_heading(line[3:], 1)
            elif line.startswith("# "):
                doc.add_heading(line[2:], 2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line:
                doc.add_paragraph(line)
        fpath = os.path.join(save_path, f"{title}.docx")
        doc.save(fpath)
        return f"Documento creado: {fpath}"

    elif action == "excel":
        try:
            import openpyxl
        except ImportError:
            return "openpyxl no instalado. Ejecutá: pip install openpyxl"
        wb = openpyxl.Workbook()
        if sheets:
            for i, s in enumerate(sheets):
                if i == 0:
                    ws = wb.active
                else:
                    ws = wb.create_sheet()
                ws.title = s.get("name", f"Sheet{i+1}")
                headers = s.get("headers", [])
                rows = s.get("rows", [])
                if headers:
                    ws.append(headers)
                for row in rows:
                    ws.append(row)
        else:
            ws = wb.active
            ws.title = "Sheet1"
            ws.append(["Dato"])
            for line in content.split("\n"):
                if line.strip():
                    ws.append([line.strip()])
        fpath = os.path.join(save_path, f"{title}.xlsx")
        wb.save(fpath)
        return f"Excel creado: {fpath}"

    elif action == "google_doc":
        try:
            from google.auth.transport.requests import Request
            from google.oauth2.credentials import Credentials
            from google_auth_oauthlib.flow import InstalledAppFlow
            from googleapiclient.discovery import build
            from pathlib import Path as P
            import json
        except ImportError:
            return "Falta google-api-python-client."
        SCOPES = ["https://www.googleapis.com/auth/documents"]
        creds_file = Path(__file__).resolve().parent.parent / "config" / "gmail_credentials.json"
        token_file = Path(__file__).resolve().parent.parent / "config" / "docs_token.json"
        creds = None
        if token_file.exists():
            creds = Credentials.from_authorized_user_file(str(token_file), SCOPES)
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(str(creds_file), SCOPES)
                creds = flow.run_local_server(port=0)
            token_file.write_text(creds.to_json())
        service = build("docs", "v1", credentials=creds)
        doc = service.documents().create(body={"title": title}).execute()
        doc_id = doc["documentId"]
        requests = []
        if content:
            requests.append({
                "insertText": {"location": {"index": 1}, "text": content}
            })
            service.documents().batchUpdate(documentId=doc_id, body={"requests": requests}).execute()
        return f"Google Doc creado: https://docs.google.com/document/d/{doc_id}"

    elif action == "google_sheet":
        try:
            from google.auth.transport.requests import Request
            from google.oauth2.credentials import Credentials
            from google_auth_oauthlib.flow import InstalledAppFlow
            from googleapiclient.discovery import build
            from pathlib import Path as P
        except ImportError:
            return "Falta google-api-python-client."
        SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]
        creds_file = Path(__file__).resolve().parent.parent / "config" / "gmail_credentials.json"
        token_file = Path(__file__).resolve().parent.parent / "config" / "sheets_token.json"
        creds = None
        if token_file.exists():
            creds = Credentials.from_authorized_user_file(str(token_file), SCOPES)
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(str(creds_file), SCOPES)
                creds = flow.run_local_server(port=0)
            token_file.write_text(creds.to_json())
        service = build("sheets", "v4", credentials=creds)
        sheet = service.spreadsheets().create(body={"properties": {"title": title}}).execute()
        sheet_id = sheet["spreadsheetId"]
        if sheets:
            for s in sheets:
                headers = s.get("headers", [])
                rows = s.get("rows", [])
                values = [headers] + rows
                service.spreadsheets().values().update(
                    spreadsheetId=sheet_id, range=f"{s.get('name','Sheet1')}!A1",
                    valueInputOption="RAW", body={"values": values}
                ).execute()
        return f"Google Sheet creado: https://docs.google.com/spreadsheets/d/{sheet_id}"

    return f"Document action '{action}' completado."
